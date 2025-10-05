import fitz
import json
import os
from bs4 import BeautifulSoup
from docx import Document


def parse_pdf(file_path: str) -> dict:
    """
    Extracts text and metadata from a PDF file.
    Returns a dictionary with cleaned content and metadata.
    """
    doc = fitz.open(file_path)

    text = ""
    for page_num, page in enumerate(doc, start=1):
        text += f"\n--- Page {page_num} ---\n"
        text += page.get_text("text")
    text = text.strip()
    metadata = {
        "title": doc.metadata.get("title", ""),
        "author": doc.metadata.get("author", ""),
        "pages": doc.page_count,
        "source": os.path.basename(file_path),
        "type": "pdf",
    }

    return {
        "doc_id": os.path.splitext(os.path.basename(file_path))[0],
        "content": text,
        "metadata": metadata,
    }


def parse_docx(file_path: str) -> dict:
    """
    Extract text and metadata from a DOCX file.
    Returns a dictionary with cleaned content and metadata.
    """
    doc = Document(file_path)

    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text.strip() + "\n"
    headings = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            headings.append(para.text.strip())
    metadata = {
        "title": headings[0] if headings else os.path.basename(file_path),
        "author": "",  # DOCX metadata optional, can add later
        "paragraphs": len(doc.paragraphs),
        "source": os.path.basename(file_path),
        "type": "docx",
    }

    return {
        "doc_id": os.path.splitext(os.path.basename(file_path))[0],
        "content": text.strip(),
        "metadata": metadata,
    }


def parse_html(file_path: str) -> dict:
    """
    Extract text and metadata from an HTML file.
    Returns a dictionary with cleaned content and metadata.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        html = f.read()

    soup = BeautifulSoup(html, "lxml")
    title = soup.title.string if soup.title else os.path.basename(file_path)
    text_parts = []

    for element in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        cleaned = element.get_text(strip=True)
        if cleaned:
            text_parts.append(cleaned)
    text = "\n".join(text_parts)
    metadata = {
        "title": title,
        "source": os.path.basename(file_path),
        "type": "html",
    }

    return {
        "doc_id": os.path.splitext(os.path.basename(file_path))[0],
        "content": text.strip(),
        "metadata": metadata,
    }


def parse_json(file_path: str) -> dict:
    """
    Extract text and metadata from a JSON file.
    Flattens the JSON into a readable text format.
    """
    with open(file_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    text_parts = []
    for key, value in data.items():
        text_parts.append(f"{key}: {value}")
    text = "\n".join(text_parts)
    metadata = {
        "title": data.get("name", os.path.basename(file_path)),
        "source": os.path.basename(file_path),
        "type": "json",
    }

    return {
        "doc_id": os.path.splitext(os.path.basename(file_path))[0],
        "content": text.strip(),
        "metadata": metadata,
    }
