from docx import Document
import os

output_dir = "documents/docx"
os.makedirs(output_dir, exist_ok=True)
file_path = os.path.join(output_dir, "messy_docx.docx")
doc = Document()
doc.add_heading("Messy DOCX Example", level=1)

# Messy spacing
doc.add_paragraph("This   DOCX    has   multiple spaces   and    line breaks.\n\n")

# Messy bullets
doc.add_paragraph("•   Item one")
doc.add_paragraph("*   Item two")
doc.add_paragraph("1.   Numbered item")

# Messy table
table = doc.add_table(rows=3, cols=3)
cells = table.rows[0].cells
cells[0].text, cells[1].text, cells[2].text = "Name", " Age ", " Country "
cells = table.rows[1].cells
cells[0].text, cells[1].text, cells[2].text = " Alice ", " 30 ", " USA "
cells = table.rows[2].cells
cells[0].text, cells[1].text, cells[2].text = " Bob ", " 25 ", " UK "

doc.save(file_path)

print("Messy DOCX generated at:", file_path)
